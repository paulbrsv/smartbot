import os
import logging
from typing import Dict, Any, List, Optional, Tuple
import re
from datetime import datetime

from app.indexing.parsers.base import BaseParser, ParsedDocument
from app.core.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


class DocxParser(BaseParser):
    """Парсер для документов Microsoft Word."""
    
    def __init__(self):
        """Инициализация парсера DOCX-документов."""
        super().__init__()
        self.supported_extensions = [".docx", ".doc"]
    
    def parse(self, file_path: str, **kwargs) -> ParsedDocument:
        """
        Разбирает DOCX-файл и возвращает его содержимое и метаданные.
        
        Args:
            file_path: Путь к файлу
            **kwargs: Дополнительные параметры
                - chunk_size: Размер чанка (по умолчанию из настроек)
                - chunk_overlap: Размер перекрытия (по умолчанию из настроек)
                - extract_headers: Извлекать ли заголовки отдельно (по умолчанию True)
                - extract_tables: Извлекать ли таблицы отдельно (по умолчанию True)
                - extract_images: Извлекать ли изображения (по умолчанию False)
                - preserve_formatting: Сохранять ли форматирование (по умолчанию True)
                - include_comments: Включать ли комментарии (по умолчанию True)
                - include_footnotes: Включать ли сноски (по умолчанию True)
            
        Returns:
            Объект ParsedDocument с содержимым и метаданными
        """
        try:
            import docx
            
            # Получаем метаданные файла
            metadata = self.get_file_metadata(file_path)
            
            # Определяем параметры
            chunk_size = kwargs.get("chunk_size", settings.processing.chunk_size)
            chunk_overlap = kwargs.get("chunk_overlap", settings.processing.chunk_overlap)
            extract_headers = kwargs.get("extract_headers", True)
            extract_tables = kwargs.get("extract_tables", True)
            extract_images = kwargs.get("extract_images", False)
            preserve_formatting = kwargs.get("preserve_formatting", True)
            include_comments = kwargs.get("include_comments", True)
            include_footnotes = kwargs.get("include_footnotes", True)
            
            # Проверяем расширение файла
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()
            
            # Открываем документ
            document = docx.Document(file_path)
            
            # Извлекаем метаданные DOCX
            docx_metadata = self._extract_docx_metadata(document)
            metadata.update(docx_metadata)
            
            # Извлекаем текст и структуру документа
            content, doc_metadata = self._extract_docx_content(
                document,
                extract_headers=extract_headers,
                extract_tables=extract_tables,
                extract_images=extract_images,
                preserve_formatting=preserve_formatting,
                include_comments=include_comments,
                include_footnotes=include_footnotes
            )
            
            # Обновляем метаданные
            metadata.update(doc_metadata)
            
            # Добавляем информацию о размере содержимого и типе
            metadata["content_length"] = len(content)
            
            if ext == ".docx":
                metadata["content_type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            else:
                metadata["content_type"] = "application/msword"
            
            # Разбиваем содержимое на чанки
            chunks = self.chunk_text(
                content,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap
            )
            
            return ParsedDocument(
                content=content,
                metadata=metadata,
                chunks=chunks
            )
        except ImportError:
            logger.error("python-docx library not installed")
            raise ImportError("python-docx is required for parsing DOCX files")
        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {e}")
            raise
    
    def _extract_docx_metadata(self, document) -> Dict[str, Any]:
        """
        Извлекает метаданные из DOCX-документа.
        
        Args:
            document: Объект DOCX-документа
            
        Returns:
            Словарь с метаданными
        """
        metadata = {}
        
        try:
            # Базовые метаданные
            core_properties = document.core_properties
            
            # Заголовок
            if core_properties.title:
                metadata["title"] = core_properties.title
            
            # Автор
            if core_properties.author:
                metadata["author"] = core_properties.author
            
            # Дата создания
            if core_properties.created:
                metadata["creation_date"] = core_properties.created.isoformat()
            
            # Дата изменения
            if core_properties.modified:
                metadata["modification_date"] = core_properties.modified.isoformat()
            
            # Дата последнего сохранения
            if core_properties.last_modified_by:
                metadata["last_modified_by"] = core_properties.last_modified_by
            
            # Редакция
            if core_properties.revision:
                metadata["revision"] = core_properties.revision
            
            # Ключевые слова
            if core_properties.keywords:
                metadata["keywords"] = core_properties.keywords
            
            # Тема
            if core_properties.subject:
                metadata["subject"] = core_properties.subject
            
            # Категория
            if core_properties.category:
                metadata["category"] = core_properties.category
            
            # Комментарии
            if core_properties.comments:
                metadata["comments"] = core_properties.comments
            
            # Статистика
            if hasattr(document, "paragraphs"):
                metadata["paragraph_count"] = len(document.paragraphs)
            
            if hasattr(document, "tables"):
                metadata["table_count"] = len(document.tables)
            
            if hasattr(document, "sections"):
                metadata["section_count"] = len(document.sections)
            
            return metadata
        except Exception as e:
            logger.warning(f"Error extracting DOCX metadata: {e}")
            return metadata
    
    def _extract_docx_content(
        self,
        document,
        extract_headers: bool = True,
        extract_tables: bool = True,
        extract_images: bool = False,
        preserve_formatting: bool = True,
        include_comments: bool = True,
        include_footnotes: bool = True
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Извлекает содержимое и структуру DOCX-документа.
        
        Args:
            document: Объект DOCX-документа
            extract_headers: Извлекать ли заголовки отдельно
            extract_tables: Извлекать ли таблицы отдельно
            extract_images: Извлекать ли изображения
            preserve_formatting: Сохранять ли форматирование
            include_comments: Включать ли комментарии
            include_footnotes: Включать ли сноски
            
        Returns:
            Кортеж (текстовое содержимое, метаданные)
        """
        text_parts = []
        metadata = {
            "headers": [],
            "table_count": 0,
            "image_count": 0,
            "footnote_count": 0,
            "comment_count": 0,
            "has_headers": False,
            "has_tables": False,
            "has_images": False,
            "has_footnotes": False,
            "has_comments": False
        }
        
        # Счетчики элементов
        header_count = 0
        table_count = 0
        image_count = 0
        
        try:
            # Обрабатываем заголовки документа
            if extract_headers:
                headers = self._extract_headers(document)
                if headers:
                    metadata["headers"] = headers
                    metadata["has_headers"] = True
                    header_count = len(headers)
            
            # Обрабатываем параграфы
            for paragraph in document.paragraphs:
                # Получаем текст параграфа
                paragraph_text = paragraph.text.strip()
                
                if not paragraph_text:
                    continue
                
                # Определяем, является ли параграф заголовком
                is_header = False
                header_level = 0
                
                if hasattr(paragraph, "style") and paragraph.style:
                    style_name = paragraph.style.name
                    
                    # Проверяем, является ли стиль заголовком
                    if re.match(r"Heading \d+|Заголовок \d+", style_name):
                        is_header = True
                        header_level = int(re.search(r"\d+", style_name).group())
                
                # Форматируем параграф
                if is_header and preserve_formatting:
                    # Добавляем символы # для заголовков
                    formatted_text = f"{'#' * header_level} {paragraph_text}"
                else:
                    formatted_text = paragraph_text
                
                # Добавляем параграф
                text_parts.append(formatted_text)
            
            # Обрабатываем таблицы
            if extract_tables:
                for table in document.tables:
                    table_count += 1
                    table_text = self._extract_table(table)
                    
                    if table_text:
                        text_parts.append(f"--- Таблица {table_count} ---")
                        text_parts.append(table_text)
                
                metadata["table_count"] = table_count
                metadata["has_tables"] = table_count > 0
            
            # Обрабатываем изображения
            if extract_images:
                # Извлечение изображений требует работы с внутренним представлением DOCX
                # Эта функциональность может быть добавлена в будущем
                # Здесь мы только указываем количество изображений
                metadata["image_count"] = image_count
                metadata["has_images"] = image_count > 0
            
            # Обрабатываем сноски (требует дополнительной работы с XML-структурой DOCX)
            if include_footnotes:
                # Упрощенная обработка
                footnote_count = 0  # Заглушка
                metadata["footnote_count"] = footnote_count
                metadata["has_footnotes"] = footnote_count > 0
            
            # Обрабатываем комментарии (требует дополнительной работы с XML-структурой DOCX)
            if include_comments:
                # Упрощенная обработка
                comment_count = 0  # Заглушка
                metadata["comment_count"] = comment_count
                metadata["has_comments"] = comment_count > 0
            
            # Объединяем текст
            content = "\n\n".join(text_parts)
            
            # Добавляем статистику
            word_count = len(re.findall(r"\b\w+\b", content))
            metadata["word_count"] = word_count
            metadata["char_count"] = len(content)
            
            # Определяем язык
            metadata["language"] = self._detect_language(content)
            
            return content, metadata
        except Exception as e:
            logger.warning(f"Error extracting DOCX content: {e}")
            # Возвращаем то, что удалось извлечь
            content = "\n\n".join(text_parts)
            return content, metadata
    
    def _extract_headers(self, document) -> List[Dict[str, Any]]:
        """
        Извлекает заголовки из документа.
        
        Args:
            document: Объект DOCX-документа
            
        Returns:
            Список словарей с информацией о заголовках
        """
        headers = []
        
        try:
            for paragraph in document.paragraphs:
                if not paragraph.text.strip():
                    continue
                
                if hasattr(paragraph, "style") and paragraph.style:
                    style_name = paragraph.style.name
                    
                    # Проверяем, является ли стиль заголовком
                    if re.match(r"Heading \d+|Заголовок \d+", style_name):
                        # Извлекаем уровень заголовка
                        level_match = re.search(r"\d+", style_name)
                        level = int(level_match.group()) if level_match else 0
                        
                        # Добавляем информацию о заголовке
                        headers.append({
                            "text": paragraph.text,
                            "level": level
                        })
            
            return headers
        except Exception as e:
            logger.warning(f"Error extracting headers: {e}")
            return headers
    
    def _extract_table(self, table) -> str:
        """
        Извлекает таблицу в текстовом формате.
        
        Args:
            table: Объект таблицы DOCX
            
        Returns:
            Текстовое представление таблицы
        """
        try:
            rows = []
            
            # Обрабатываем строки таблицы
            for row in table.rows:
                cells = []
                
                # Обрабатываем ячейки строки
                for cell in row.cells:
                    # Извлекаем текст ячейки
                    cell_text = cell.text.strip().replace("\n", " ")
                    cells.append(cell_text)
                
                # Добавляем строку в виде текста с разделителями
                row_text = " | ".join(cells)
                rows.append(row_text)
            
            # Объединяем строки
            table_text = "\n".join(rows)
            
            return table_text
        except Exception as e:
            logger.warning(f"Error extracting table: {e}")
            return ""
    
    def _detect_language(self, text: str) -> str:
        """
        Определяет язык текста (простая эвристика).
        
        Args:
            text: Текст для анализа
            
        Returns:
            Код языка (en, ru, и т.д.)
        """
        try:
            # Берем небольшой фрагмент текста для анализа
            sample = text[:2000].lower()
            
            # Словарь с частотными словами для разных языков
            language_words = {
                "ru": {"и", "в", "не", "на", "я", "что", "с", "по", "это", "для", "как", "к", "из", "а", "то"},
                "en": {"the", "and", "to", "of", "a", "in", "is", "that", "for", "it", "with", "as", "on", "be", "at"},
                "de": {"der", "die", "und", "in", "den", "von", "zu", "das", "mit", "sich", "des", "auf", "für", "ist", "im"},
                "fr": {"le", "la", "et", "les", "des", "en", "un", "une", "du", "dans", "est", "que", "pour", "qui", "sur"}
            }
            
            # Подсчитываем количество совпадений для каждого языка
            matches = {}
            words = re.findall(r"\b\w+\b", sample)
            
            for lang, lang_words in language_words.items():
                matches[lang] = sum(1 for word in words if word in lang_words)
            
            # Определяем язык с наибольшим количеством совпадений
            if not matches:
                return "unknown"
            
            best_lang = max(matches, key=matches.get)
            
            # Если количество совпадений слишком мало, считаем язык неопределенным
            if matches[best_lang] < 3:
                return "unknown"
            
            return best_lang
        except Exception as e:
            logger.warning(f"Error detecting language: {e}")
            return "unknown"


# Регистрируем парсер в фабрике
from app.indexing.parsers.base import ParserFactory

ParserFactory.register_parser(DocxParser)